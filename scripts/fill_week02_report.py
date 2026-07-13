#!/usr/bin/env python3
"""
fill_week02_report.py
Mo file DOCX mau, dien noi dung Tuan 02, luu file dau ra.

Su dung: python scripts/fill_week02_report.py

SOURCE: Mau_BaoCao_Tuan02_BanThaoTichLuy_TheoChuong_50DeTai.docx
OUTPUT: LuongNguyenNgocDinh_baocaotieuluanlan01.docx
"""
import sys, os, io, hashlib
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from datetime import date

# ── Paths ──────────────────────────────────────────────
SOURCE = r'C:\Users\hqhuy\Downloads\Mau_BaoCao_Tuan02_BanThaoTichLuy_TheoChuong_50DeTai.docx'
OUTPUT = r'C:\Users\hqhuy\Downloads\LuongNguyenNgocDinh_baocaotieuluanlan01.docx'
REPO_OUTPUT = r'C:\Users\hqhuy\iot-gateway-security-lab\report\LuongNguyenNgocDinh_baocaotieuluanlan01.docx'

# ── Data ───────────────────────────────────────────────
TODAY = date.today().strftime('%d/%m/%Y')
STUDENT = 'Lương Nguyễn Ngọc Đinh'
MSSV = '231A010001'
TOPIC_CODE = '04'
TOPIC_NAME = 'VAI TRÒ CỦA GATEWAY TRONG BẢO MẬT IOT'
DIRECTION = 'A'
CLASS_NAME = 'INT4410 - Bảo mật IoT'
GITHUB = 'https://github.com/dinhvaren/iot-gateway-security-lab'
COMMIT_ID = '407f53c'

# ── Helpers ────────────────────────────────────────────

def clear_fill(p, text):
    """Xoa toan bo runs cua paragraph va dat text moi vao run dau tien."""
    runs = p.runs
    for r in runs:
        r.text = ''
    if runs:
        runs[0].text = text
    else:
        p.add_run(text)

def fill_cell(cell, text):
    """Dat text cho cell, xoa noi dung cu."""
    for para in cell.paragraphs:
        for r in para.runs:
            r.text = ''
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

def sha256(path):
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        h.update(f.read())
    return h.hexdigest().upper()

# ── Main ───────────────────────────────────────────────
print("=" * 60)
print("fill_week02_report.py")
print(f"Source: {SOURCE}")
print(f"Output: {OUTPUT}")
print("=" * 60)

# 1. Tinh SHA256 file nguon
src_hash = sha256(SOURCE)
print(f"\nSOURCE SHA256: {src_hash}")

# 2. Open template
print("\nOpening template...")
doc = Document(SOURCE)
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")

# ── COVER PAGE ─────────────────────────────────────────
print("\n--- Cover Page ---")

# p[5] = placeholder ten de tai
clear_fill(doc.paragraphs[5], TOPIC_NAME)
print(f"  p[5] <- {TOPIC_NAME}")

# Table 0: Cover info
t0 = doc.tables[0]
fill_cell(t0.rows[0].cells[1], TOPIC_CODE)       # Ma de tai
fill_cell(t0.rows[1].cells[1], DIRECTION)          # Huong
fill_cell(t0.rows[2].cells[1], STUDENT)            # Sinh vien
fill_cell(t0.rows[3].cells[1], MSSV)               # MSSV
fill_cell(t0.rows[4].cells[1], CLASS_NAME)         # Lop
fill_cell(t0.rows[5].cells[1], '')                 # Giang vien (trong)
fill_cell(t0.rows[6].cells[1], GITHUB)             # GitHub
fill_cell(t0.rows[7].cells[1], TODAY)              # Ngay nop
print("  Table 0 (cover): filled 8 rows")

# ── PART A: QUICK CHECK (Table 3) ─────────────────────
print("\n--- Part A: Quick Check ---")
t3 = doc.tables[3]
fill_cell(t3.rows[1].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[1].cells[3], f'Repo: {GITHUB} | README.md | docs/')
fill_cell(t3.rows[2].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[2].cells[3], 'Khoảng 2–3 trang (xem Chương 1)')
fill_cell(t3.rows[3].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[3].cells[3], 'Khoảng 2–3 trang bản nháp (xem Chương 2)')
fill_cell(t3.rows[4].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[4].cells[3], 'Mục 2.4 – 5 nguồn chính')
fill_cell(t3.rows[5].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[5].cells[3], f'{GITHUB} | Commit {COMMIT_ID}')
fill_cell(t3.rows[6].cells[2], '☒ Đạt  ☐ Bổ sung')
fill_cell(t3.rows[6].cells[3], 'Phần C – Dàn ý Chương 3 + docs/KIEN_TRUC_HE_THONG.md')
print("  Table 3 (quick check): filled 6 rows")

# ── CHAPTER 1 (100%) ──────────────────────────────────
print("\n--- Chapter 1 ---")

# 1.1 Bối cảnh và lý do chọn đề tài (p[23], p[24] dots)
s1_1 = (
    "Trong những năm gần đây, Internet of Things (IoT) đã phát triển nhanh chóng và được triển khai rộng rãi "
    "trong nhiều lĩnh vực như nhà máy thông minh, nông nghiệp chính xác, y tế từ xa và giám sát môi trường. "
    "Trong các hệ thống này, thiết bị cảm biến thường đảm nhiệm vai trò thu thập dữ liệu từ môi trường vật lý "
    "và chuyển tiếp dữ liệu tới hệ thống backend hoặc nền tảng đám mây để xử lý và ra quyết định [1].\n\n"
    "Tuy nhiên, phần lớn thiết bị IoT bị hạn chế về tài nguyên xử lý, dung lượng lưu trữ và khả năng triển khai "
    "các cơ chế bảo mật phức tạp. Nếu dữ liệu được gửi trực tiếp từ thiết bị cảm biến lên cloud mà không qua "
    "một lớp kiểm soát trung gian, hệ thống có nguy cơ tiếp nhận dữ liệu từ thiết bị không được phép, token "
    "xác thực không hợp lệ, payload thiếu trường hoặc sai cấu trúc, giá trị cảm biến vượt ngoài ngưỡng cho phép, "
    "dữ liệu cũ (stale data) hoặc bị gửi lặp lại (replay) [2].\n\n"
    "Gateway – thiết bị trung gian nằm giữa lớp cảm biến và hạ tầng đám mây – có thể đảm nhiệm vai trò kiểm soát "
    "và lọc dữ liệu trước khi dữ liệu được chuyển tiếp lên các hệ thống phía trên. Thay vì chỉ đơn thuần là bộ "
    "định tuyến, gateway có thể được thiết kế như một điểm thực thi chính sách bảo mật, giúp giảm tải cho thiết "
    "bị biên vốn bị hạn chế tài nguyên [1].\n\n"
    "Xuất phát từ bối cảnh đó, đề tài này tập trung nghiên cứu và xây dựng một mô hình gateway an ninh (Security "
    "Gateway) dựa trên nền tảng Node-RED, chạy trong môi trường Docker. Mô hình hướng tới việc minh họa cách "
    "gateway có thể thực hiện các chức năng xác thực thiết bị, kiểm tra cấu trúc dữ liệu, lọc giá trị bất thường, "
    "phát hiện dữ liệu cũ hoặc lặp và ghi nhận toàn bộ quyết định dưới dạng audit log. Các tình huống và kết quả "
    "thử nghiệm cụ thể sẽ được trình bày trong các chương tiếp theo của báo cáo."
)
clear_fill(doc.paragraphs[23], s1_1)
clear_fill(doc.paragraphs[24], '')
print("  1.1 done")

# 1.2 Phát biểu vấn đề bảo mật (p[26], p[27])
s1_2 = (
    "Tài sản cần bảo vệ trong mô hình bao gồm: dữ liệu cảm biến (nhiệt độ, độ ẩm, trạng thái thiết bị...), "
    "danh tính thiết bị (device identity), token hoặc thông tin xác thực, luồng dữ liệu gửi lên cloud và "
    "audit log ghi nhận toàn bộ quyết định tại gateway.\n\n"
    "Các mối đe dọa chính được xác định bao gồm: giả mạo nguồn gửi (device spoofing), gửi payload sai cấu "
    "trúc hoặc thiếu trường (malformed payload), dữ liệu cảm biến vượt ngưỡng vật lý cho phép (anomalous data), "
    "gửi lại dữ liệu cũ (stale data), tấn công phát lại (replay attack) và thiếu khả năng giám sát, truy vết "
    "khi xảy ra sự cố.\n\n"
    "Nếu những mối đe dọa này không được kiểm soát, hậu quả có thể bao gồm: hệ thống cloud xử lý dữ liệu sai "
    "dẫn đến quyết định sai lệch, lãng phí tài nguyên xử lý, khó khăn trong điều tra sự cố và mất tính toàn "
    "vẹn của dữ liệu.\n\n"
    "Khoảng trống mà đề tài hướng tới xử lý là: xây dựng một mô hình gateway đóng vai trò lớp kiểm soát dữ "
    "liệu trung gian, thực hiện các bước kiểm tra trước khi dữ liệu được chuyển tiếp lên cloud, đồng thời "
    "ghi nhận toàn bộ quyết định để phục vụ truy vết và đánh giá."
)
clear_fill(doc.paragraphs[26], s1_2)
clear_fill(doc.paragraphs[27], '')
print("  1.2 done")

# 1.3 Mục tiêu (p[29], p[30])
s1_3 = (
    "MT-01: Phân tích vai trò của gateway trong kiến trúc và bảo mật IoT.\n"
    "Đầu ra dự kiến: Mô tả kiến trúc hệ thống, luồng dữ liệu và các điểm kiểm soát tại gateway; "
    "sơ đồ tổng quan được trình bày tại Chương 2.\n\n"
    "MT-02: Thiết kế mô hình gateway bảo mật bằng Node-RED, tiếp nhận và xử lý dữ liệu IoT giả lập.\n"
    "Đầu ra dự kiến: Sơ đồ kiến trúc chi tiết và flow Node-RED dự kiến; sẽ triển khai cụ thể "
    "trong Chương 3 và Chương 4.\n\n"
    "MT-03: Đề xuất và triển khai các cơ chế kiểm tra tại gateway, bao gồm: xác thực thiết bị, "
    "kiểm tra cấu trúc payload, lọc dữ liệu bất thường và ghi nhận audit log.\n"
    "Đầu ra dự kiến: Danh sách rule/checklist, cấu hình gateway và mẫu audit log ACCEPT/REJECT; "
    "các tuần tiếp theo sẽ hoàn thiện và kiểm thử.\n\n"
    "MT-04: Xây dựng kế hoạch đánh giá mô hình bằng tập payload hợp lệ và không hợp lệ.\n"
    "Đầu ra dự kiến: Test plan, bảng test case, log kiểm thử, ảnh chụp màn hình và video demo "
    "trong các tuần tiếp theo."
)
clear_fill(doc.paragraphs[29], s1_3)
clear_fill(doc.paragraphs[30], '')
print("  1.3 done")

# 1.4 Đối tượng, phạm vi (p[32], p[33])
s1_4 = (
    "Đối tượng nghiên cứu của đề tài bao gồm: dữ liệu cảm biến IoT giả lập, gateway (Node-RED Security "
    "Gateway), luồng chuyển dữ liệu lên cloud sink giả lập, nền tảng Node-RED và môi trường Docker.\n\n"
    "Phạm vi thực hiện: Triển khai hoàn toàn trên môi trường local sử dụng Docker; không tấn công hoặc "
    "can thiệp vào hệ thống thật; chỉ sử dụng dữ liệu cảm biến giả lập và tài nguyên thuộc quyền sở hữu "
    "của người thực hiện. Trọng tâm kỹ thuật tập trung vào bốn nhóm chức năng tại gateway: xác thực "
    "(authentication), kiểm tra cấu trúc dữ liệu (validation), lọc dữ liệu bất thường (filtering) và ghi "
    "nhận nhật ký (audit logging).\n\n"
    "Các nội dung không thực hiện trong giai đoạn hiện tại bao gồm: triển khai trên thiết bị IoT vật lý, "
    "hệ thống production, hạ tầng khóa công khai (PKI) hoàn chỉnh, xác thực hai chiều mTLS hoàn chỉnh, "
    "hệ thống SIEM, kiểm thử tải quy mô lớn và áp dụng machine learning để phát hiện bất thường."
)
clear_fill(doc.paragraphs[32], s1_4)
clear_fill(doc.paragraphs[33], '')
print("  1.4 done")

# 1.5 Sản phẩm dự kiến (p[35], p[36])
s1_5 = (
    "Các sản phẩm và kết quả dự kiến của đề tài bao gồm:\n\n"
    "- Repository GitHub chứa toàn bộ mã nguồn, cấu hình và tài liệu.\n"
    "- README hướng dẫn cài đặt, cấu hình và vận hành mô hình.\n"
    "- File Node-RED flow (flows.json) mô tả luồng xử lý tại gateway.\n"
    "- File Docker Compose (docker-compose.yml) để triển khai môi trường.\n"
    "- Tập payload IoT giả lập dùng cho kiểm thử.\n"
    "- Bộ rule/checklist kiểm tra tại gateway (xác thực, validation, filtering, logging).\n"
    "- Mẫu audit log ghi nhận quyết định ACCEPT/REJECT kèm nguyên nhân.\n"
    "- Bảng test case đối chiếu payload hợp lệ và không hợp lệ.\n"
    "- Bảng phân tích rủi ro và biện pháp giảm thiểu.\n"
    "- Ảnh chụp màn hình và video demo minh họa kết quả.\n"
    "- Báo cáo tích lũy hoàn chỉnh vào Buổi 06.\n\n"
    "Các sản phẩm trên sẽ được hoàn thiện dần qua từng tuần và đóng gói lần cuối ở Buổi 06."
)
clear_fill(doc.paragraphs[35], s1_5)
clear_fill(doc.paragraphs[36], '')
print("  1.5 done")

# 1.6 Cấu trúc báo cáo (p[38], p[39])
s1_6 = (
    "Báo cáo được tổ chức thành sáu chương. Chương 1 trình bày tổng quan đề tài, bao gồm bối cảnh, "
    "vấn đề bảo mật, mục tiêu, phạm vi và sản phẩm dự kiến. Chương 2 cung cấp cơ sở lý thuyết về "
    "kiến trúc IoT, các khái niệm bảo mật liên quan và tổng quan các công cụ, chuẩn và công trình "
    "liên quan. Chương 3 mô tả phương pháp thực hiện, mô hình kiến trúc đề xuất và môi trường triển "
    "khai. Chương 4 trình bày quá trình triển khai và kết quả thử nghiệm. Chương 5 thực hiện đánh giá "
    "bảo mật dựa trên mô hình đã triển khai. Chương 6 tổng kết, nêu hạn chế và đề xuất hướng phát triển."
)
clear_fill(doc.paragraphs[38], s1_6)
clear_fill(doc.paragraphs[39], '')
print("  1.6 done")

# ── TABLE 5: MỤC TIÊU - ĐẦU RA ────────────────────────
print("\n--- Target-Output Table ---")
t5 = doc.tables[5]
fill_cell(t5.rows[1].cells[1], 'Sơ đồ kiến trúc và phân tích vai trò gateway trong bảo mật IoT.')
fill_cell(t5.rows[1].cells[2], 'Đối chiếu sơ đồ, mô tả luồng dữ liệu và tài liệu tham khảo.')
fill_cell(t5.rows[1].cells[3], 'Chương 1, Chương 2')
fill_cell(t5.rows[2].cells[1], 'Mô hình Node-RED Security Gateway tiếp nhận và xử lý dữ liệu IoT giả lập.')
fill_cell(t5.rows[2].cells[2], 'Flow Node-RED, cấu hình Docker và ảnh giao diện trong các tuần triển khai.')
fill_cell(t5.rows[2].cells[3], 'Chương 3, Chương 4')
fill_cell(t5.rows[3].cells[1], 'Bộ rule/checklist xác thực, kiểm tra cấu trúc, lọc dữ liệu và audit log.')
fill_cell(t5.rows[3].cells[2], 'Code/config, log ACCEPT/REJECT và test case đối chiếu.')
fill_cell(t5.rows[3].cells[3], 'Chương 3, Chương 4, Chương 5')
fill_cell(t5.rows[4].cells[1], 'Kế hoạch và kết quả kiểm thử mô hình bằng payload hợp lệ và không hợp lệ.')
fill_cell(t5.rows[4].cells[2], 'Bảng test case, screenshot, log kiểm thử và video demo.')
fill_cell(t5.rows[4].cells[3], 'Chương 4, Chương 5')
print("  Table 5: filled 4 rows")

# ── CHAPTER 2 (50-70%) ────────────────────────────────
print("\n--- Chapter 2 ---")

# 2.1 Kiến trúc (p[45], p[46], p[47])
s2_1 = (
    "Mô hình kiến trúc dự kiến của đề tài được tổ chức thành ba vùng chính: vùng thiết bị/dữ liệu đầu vào, "
    "vùng gateway (ranh giới tin cậy) và vùng cloud/backend giả lập.\n\n"
    "Vùng thiết bị/dữ liệu đầu vào: Bao gồm các IoT Device Simulator – chương trình giả lập thiết bị cảm biến, "
    "có nhiệm vụ sinh dữ liệu (nhiệt độ, độ ẩm, trạng thái...) và gửi đến gateway qua giao thức HTTP hoặc "
    "MQTT. Dữ liệu tại vùng này được coi là chưa được tin cậy (untrusted).\n\n"
    "Vùng gateway – Node-RED Security Gateway: Đây là thành phần trung tâm của mô hình, hoạt động như một "
    "điểm thực thi chính sách bảo mật (Policy Enforcement Point). Tại đây, dữ liệu đi qua một chuỗi các bước "
    "kiểm tra tuần tự: (1) Authentication – xác thực danh tính thiết bị qua token hoặc khóa API; (2) Device/Data "
    "Validation – kiểm tra cấu trúc payload, các trường bắt buộc và kiểu dữ liệu; (3) Filtering – lọc giá trị "
    "cảm biến theo ngưỡng cho phép; (4) Freshness/Replay Check – kiểm tra timestamp và phát hiện dữ liệu cũ "
    "hoặc gửi lặp; (5) Audit Logging – ghi nhận toàn bộ quyết định ACCEPT/REJECT kèm nguyên nhân.\n\n"
    "Vùng cloud/backend giả lập (Cloud Sink Simulator): Nhận dữ liệu đã được gateway kiểm tra và chấp nhận. "
    "Trong phạm vi đề tài, đây là một endpoint HTTP đơn giản ghi nhận dữ liệu đến, phục vụ mục đích minh họa "
    "và kiểm thử.\n\n"
    "Sơ đồ kiến trúc dự kiến:\n\n"
    "IoT Device Simulator --> [HTTP/MQTT] --> Node-RED Security Gateway --> [HTTP] --> Cloud Sink Simulator\n"
    "                                            |\n"
    "                                            +-- Authentication (token/API key)\n"
    "                                            +-- Device/Data Validation (schema, required fields)\n"
    "                                            +-- Filtering (threshold, range check)\n"
    "                                            +-- Freshness/Replay Check (timestamp, nonce)\n"
    "                                            +-- Audit Logging (ACCEPT/REJECT + reason)\n\n"
    "Hình 2.1. Kiến trúc dự kiến của mô hình IoT Security Gateway\n\n"
    "Ranh giới tin cậy (trust boundary) được xác định tại gateway: mọi dữ liệu từ vùng thiết bị được coi là "
    "không tin cậy cho đến khi vượt qua toàn bộ các bước kiểm tra. Chỉ dữ liệu đã được gateway chấp nhận mới "
    "được phép chuyển tiếp vào vùng cloud."
)
clear_fill(doc.paragraphs[45], s2_1)
clear_fill(doc.paragraphs[46], '')
clear_fill(doc.paragraphs[47], '')
print("  2.1 done")

# 2.2 Khái niệm bảo mật (p[49], p[50], p[51])
s2_2 = (
    "Các khái niệm bảo mật sau đây được sử dụng trực tiếp trong thiết kế và đánh giá mô hình gateway:\n\n"
    "Tài sản (Asset): Trong mô hình, tài sản bao gồm dữ liệu cảm biến, danh tính thiết bị, token xác thực "
    "và audit log. Gateway có trách nhiệm bảo vệ tính toàn vẹn của dữ liệu trước khi dữ liệu được chuyển "
    "tiếp lên cloud.\n\n"
    "Mối đe dọa (Threat): Các hành vi có thể gây tổn hại đến tài sản, như giả mạo thiết bị, gửi dữ liệu sai "
    "cấu trúc, gửi dữ liệu vượt ngưỡng, phát lại gói tin cũ. Gateway được thiết kế để phát hiện và chặn các "
    "mối đe dọa này tại biên.\n\n"
    "Lỗ hổng (Vulnerability): Điểm yếu trong kiến trúc có thể bị khai thác, ví dụ thiếu cơ chế xác thực giữa "
    "thiết bị và cloud, hoặc không có bước kiểm tra cấu trúc payload. Mô hình đề xuất sẽ bổ sung các cơ chế "
    "này tại gateway.\n\n"
    "Rủi ro (Risk): Khả năng mối đe dọa khai thác lỗ hổng gây tác động tiêu cực. Bảng phân tích rủi ro sẽ "
    "được trình bày chi tiết tại Chương 5.\n\n"
    "Tính bí mật (Confidentiality), Tính toàn vẹn (Integrity), Tính sẵn sàng (Availability) – bộ ba CIA: "
    "Gateway góp phần bảo đảm tính toàn vẹn dữ liệu (qua validation và filtering) và tính sẵn sàng (qua việc "
    "chặn dữ liệu không hợp lệ trước khi tiêu tốn tài nguyên cloud).\n\n"
    "Xác thực (Authentication): Gateway thực hiện xác thực thiết bị thông qua token hoặc API key trước khi "
    "chấp nhận dữ liệu.\n\n"
    "Validation (Kiểm tra cấu trúc): Gateway kiểm tra cấu trúc payload, đảm bảo đầy đủ các trường bắt buộc "
    "và đúng kiểu dữ liệu.\n\n"
    "Filtering (Lọc dữ liệu): Gateway so sánh giá trị cảm biến với ngưỡng cho phép và từ chối dữ liệu bất thường.\n\n"
    "Freshness (Tính mới của dữ liệu): Gateway kiểm tra timestamp để đảm bảo dữ liệu không quá cũ, tránh stale data.\n\n"
    "Replay Protection (Chống phát lại): Gateway phát hiện và từ chối các gói tin bị gửi lặp dựa trên nonce "
    "hoặc timestamp kết hợp với device ID.\n\n"
    "Audit Logging (Ghi nhật ký): Toàn bộ quyết định ACCEPT/REJECT tại gateway đều được ghi lại kèm nguyên nhân, "
    "phục vụ truy vết và đánh giá sau này."
)
clear_fill(doc.paragraphs[49], s2_2)
clear_fill(doc.paragraphs[50], '')
clear_fill(doc.paragraphs[51], '')
print("  2.2 done")

# 2.3 Nhóm đề tài - Table 7 (đánh dấu A)
t7 = doc.tables[7]
fill_cell(t7.rows[1].cells[0], '☒ A')
print("  2.3 done (direction A marked in Table 7)")

# 2.4 Nguồn tham khảo - Table 8
t8 = doc.tables[8]
sources = [
    ('1', 'Node-RED', 'https://github.com/node-red/node-red',
     'Nền tảng flow-based programming để xây dựng mô hình gateway xử lý dữ liệu IoT theo luồng.'),
    ('2', 'Node-RED Docker', 'https://github.com/node-red/node-red-docker',
     'Triển khai Node-RED trong môi trường Docker để tái lập và cô lập mô hình thử nghiệm.'),
    ('3', 'OWASP IoT Security Verification Standard (ISVS)',
     'https://github.com/OWASP/IoT-Security-Verification-Standard-ISVS',
     'Tham khảo yêu cầu và nguyên tắc đánh giá bảo mật IoT; sử dụng làm khung tham chiếu cho thiết kế rule/checklist tại gateway.'),
    ('4', 'NISTIR 8259A – IoT Device Cybersecurity Capability Core Baseline',
     'https://nvlpubs.nist.gov/nistpubs/ir/2020/NIST.IR.8259A.pdf',
     'Tham khảo năng lực bảo mật cơ bản cho thiết bị IoT; tài liệu hiện hành tại ngày truy cập.'),
    ('5', 'ETSI EN 303 645 – Cyber Security for Consumer Internet of Things',
     'https://www.etsi.org/deliver/etsi_en/303600_303699/303645/02.01.01_60/en_303645v020101p.pdf',
     'Tham khảo nguyên tắc bảo mật cho thiết bị IoT tiêu dùng; tài liệu hiện hành tại ngày truy cập.'),
]
for i, (stt, name, url, desc) in enumerate(sources):
    fill_cell(t8.rows[i+1].cells[0], stt)
    fill_cell(t8.rows[i+1].cells[1], name)
    fill_cell(t8.rows[i+1].cells[2], url)
    fill_cell(t8.rows[i+1].cells[3], desc)
    fill_cell(t8.rows[i+1].cells[4], TODAY)
print("  2.4 done (5 sources)")

# 2.5 Công trình liên quan (p[55], p[56], p[57])
s2_5 = (
    "Node-RED là một nền tảng lập trình trực quan dạng flow-based, được phát triển bởi JS Foundation. "
    "Ưu điểm chính của Node-RED là khả năng mô hình hóa luồng dữ liệu một cách trực quan thông qua giao "
    "diện kéo-thả, tích hợp sẵn nhiều node hỗ trợ các giao thức phổ biến trong IoT như MQTT, HTTP, "
    "WebSocket. Điều này khiến Node-RED trở thành lựa chọn phù hợp để xây dựng nguyên mẫu gateway trong "
    "phạm vi đề tài. Tuy nhiên, bản thân Node-RED không tự động biến mọi flow thành gateway an toàn – "
    "các chính sách bảo mật (authentication, validation, filtering) cần được thiết kế và cấu hình đúng "
    "bởi người triển khai.\n\n"
    "Node-RED Docker cung cấp khả năng đóng gói và tái lập môi trường thử nghiệm một cách nhất quán. "
    "Ưu điểm của Docker là giúp cô lập môi trường, dễ dàng triển khai lại trên các máy khác nhau mà không "
    "phụ thuộc vào cấu hình hệ điều hành. Tuy nhiên, Docker chỉ cung cấp môi trường đóng gói; nó không "
    "thay thế logic bảo mật bên trong gateway.\n\n"
    "OWASP IoT Security Verification Standard (ISVS) cung cấp một khung tham chiếu toàn diện về các yêu "
    "cầu bảo mật cho hệ thống IoT, bao gồm xác thực, mã hóa, quản lý lỗ hổng và ghi nhận sự kiện. Ưu "
    "điểm của ISVS là có cấu trúc rõ ràng, dễ ánh xạ sang các yêu cầu cụ thể. Hạn chế là ISVS là tiêu "
    "chuẩn/khung yêu cầu, cần được chuyển đổi thành rule, checklist và test case phù hợp với mô hình "
    "cụ thể của đề tài.\n\n"
    "Phần đề tài kế thừa: Mô hình đề xuất sẽ kết hợp Node-RED (nền tảng flow-based), Docker (môi trường "
    "đóng gói) và các nguyên tắc từ OWASP ISVS, NISTIR 8259A và ETSI EN 303 645 để xây dựng một gateway "
    "mô phỏng có khả năng thực hiện các chức năng kiểm soát bảo mật cơ bản. Việc đánh giá thực nghiệm "
    "đầy đủ sẽ được thực hiện trong các tuần tiếp theo."
)
clear_fill(doc.paragraphs[55], s2_5)
clear_fill(doc.paragraphs[56], '')
clear_fill(doc.paragraphs[57], '')
print("  2.5 done")

# 2.6 Tiểu kết (p[59])
s2_6 = (
    "Chương 2 đã xác định kiến trúc dự kiến của mô hình IoT Security Gateway với ba vùng chính (thiết bị, "
    "gateway, cloud), trong đó gateway được xem là ranh giới tin cậy và là điểm thực thi chính sách bảo mật. "
    "Các khái niệm bảo mật cốt lõi như xác thực, validation, filtering, freshness, replay protection và "
    "audit logging đã được làm rõ và gắn với vai trò cụ thể của gateway. Năm nguồn tham khảo chính – bao "
    "gồm Node-RED, Node-RED Docker, OWASP ISVS, NISTIR 8259A và ETSI EN 303 645 – đã được xác định và mô "
    "tả phần sử dụng. Các kiến thức nền này sẽ làm cơ sở cho việc thiết kế phương pháp và mô hình chi tiết "
    "trong Chương 3. Chương 2 sẽ tiếp tục được hoàn thiện ở Buổi 03 sau khi nhận góp ý từ giảng viên."
)
clear_fill(doc.paragraphs[59], s2_6)
print("  2.6 done")

# ── CHAPTER 3-6 PLACEHOLDERS ──────────────────────────
print("\n--- Chapters 3-6 (placeholders) ---")
placeholder_map = {
    64: 'Nội dung dự kiến hoàn thiện ở Buổi 03. Chương này sẽ trình bày phương pháp thực hiện, mô hình kiến trúc đề xuất chi tiết, môi trường triển khai (Docker + Node-RED), quy trình triển khai và tiêu chí đánh giá mô hình gateway.',
    66: 'Nội dung dự kiến hoàn thiện ở Buổi 03.',
    68: 'Nội dung dự kiến hoàn thiện ở Buổi 03.',
    70: 'Nội dung dự kiến hoàn thiện ở Buổi 03.',
    72: 'Nội dung dự kiến hoàn thiện ở Buổi 03.',
    74: 'Nội dung dự kiến hoàn thiện ở Buổi 03.',
    76: 'Nội dung dự kiến hoàn thiện ở Buổi 04. Chương này sẽ trình bày quá trình triển khai mô hình, các thành phần sản phẩm, kịch bản thử nghiệm và kết quả đạt được.',
    78: 'Nội dung dự kiến hoàn thiện ở Buổi 04.',
    80: 'Nội dung dự kiến hoàn thiện ở Buổi 04.',
    82: 'Nội dung dự kiến hoàn thiện ở Buổi 04.',
    84: 'Nội dung dự kiến hoàn thiện ở Buổi 04.',
    86: 'Nội dung dự kiến hoàn thiện ở Buổi 04.',
    88: 'Nội dung dự kiến hoàn thiện ở Buổi 05. Chương này sẽ thực hiện đánh giá bảo mật mô hình, bao gồm phân tích tài sản, mối đe dọa, lỗ hổng, ma trận rủi ro và biện pháp giảm thiểu.',
    90: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    92: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    94: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    96: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    98: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    100: 'Nội dung dự kiến hoàn thiện ở Buổi 05. Chương này sẽ tổng kết kết quả đạt được, nêu hạn chế và đề xuất hướng phát triển trong tương lai.',
    102: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    104: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
    106: 'Nội dung dự kiến hoàn thiện ở Buổi 05.',
}
for idx, text in placeholder_map.items():
    clear_fill(doc.paragraphs[idx], text)
print(f"  Filled {len(placeholder_map)} placeholder paragraphs")

# ── APPENDIX A: GÓP Ý (Table 11) ──────────────────────
print("\n--- Appendix A ---")
t11 = doc.tables[11]
fill_cell(t11.rows[1].cells[0], 'Cần cập nhật minh chứng (repo, README, đề cương, danh mục tài liệu tham khảo).')
fill_cell(t11.rows[1].cells[1], 'Đã bổ sung link repository GitHub, file README.md, đề cương Tuần 02 (docs/DE_CUONG_TUAN_02.md) và danh mục tối thiểu 5 nguồn tham khảo tại Mục 2.4.')
fill_cell(t11.rows[1].cells[2], 'Trang bìa, Mục 2.4 và Phụ lục B.')
fill_cell(t11.rows[1].cells[3], 'Đã xử lý')
print("  Table 11: filled")

# ── APPENDIX B: MINH CHỨNG (Table 12) ─────────────────
print("\n--- Appendix B ---")
t12 = doc.tables[12]
evidences = [
    ('MC-01', 'Repository', GITHUB, 'Đã tạo repository làm việc cho đề tài, bao gồm Docker Compose, Node-RED flow, tài liệu và minh chứng.'),
    ('MC-02', 'README', f'{GITHUB}/blob/main/README.md', 'README mô tả đề tài, mục tiêu, phạm vi, kiến trúc và kế hoạch thực hiện.'),
    ('MC-03', 'Đề cương', f'{GITHUB}/blob/main/docs/DE_CUONG_TUAN_02.md', 'Đề cương chi tiết bao gồm mục tiêu, phạm vi, kế hoạch thực hiện và sản phẩm dự kiến.'),
    ('MC-04', 'Tài liệu tham khảo', f'{GITHUB}/blob/main/docs/TAI_LIEU_THAM_KHAO.md', 'Danh mục 5 nguồn chính kèm URL, mô tả phần sử dụng và ngày truy cập.'),
    ('MC-05', 'Kế hoạch/kiến trúc', f'{GITHUB}/blob/main/docs/KIEN_TRUC_HE_THONG.md', 'Tài liệu mô tả kiến trúc hệ thống dự kiến và kế hoạch hoàn thiện Chương 2, chuẩn bị Chương 3.'),
]
for i, (mid, mtype, mlink, mdesc) in enumerate(evidences):
    fill_cell(t12.rows[i+1].cells[0], mid)
    fill_cell(t12.rows[i+1].cells[1], mtype)
    fill_cell(t12.rows[i+1].cells[2], mlink)
    fill_cell(t12.rows[i+1].cells[3], mdesc)
    fill_cell(t12.rows[i+1].cells[4], TODAY)
print("  Table 12: filled 5 rows")

# ── APPENDIX C: CAM KẾT (Table 13) ────────────────────
print("\n--- Appendix C ---")
t13 = doc.tables[13]
fill_cell(t13.rows[0].cells[1], 'Hoàn thiện Chương 2 sau góp ý của giảng viên và viết Chương 3 về phương pháp, mô hình kiến trúc và thiết kế chi tiết.')
fill_cell(t13.rows[1].cells[1], 'Sơ đồ kiến trúc hoàn chỉnh; flow Node-RED prototype; môi trường Docker chạy thử; test plan/checklist bản đầu; ảnh chụp và debug log ban đầu.')
fill_cell(t13.rows[2].cells[1], 'Buổi 03 (theo lịch học phần)')
fill_cell(t13.rows[3].cells[1], 'Góp ý về phạm vi các luật kiểm tra tại gateway và mức độ chi tiết của mô hình thử nghiệm.')
print("  Table 13: filled")

# ── SAVE ──────────────────────────────────────────────
print("\n--- Saving ---")
doc.save(OUTPUT)
print(f"  Saved: {OUTPUT}")
print(f"  Size: {os.path.getsize(OUTPUT)} bytes")

# Also save to repo
os.makedirs(os.path.dirname(REPO_OUTPUT), exist_ok=True)
doc.save(REPO_OUTPUT)
print(f"  Saved: {REPO_OUTPUT}")

# ── VERIFY ────────────────────────────────────────────
print("\n--- Verify ---")
out_hash = sha256(OUTPUT)
print(f"SOURCE SHA256: {src_hash}")
print(f"OUTPUT SHA256: {out_hash}")
print(f"HASHES DIFFER: {src_hash != out_hash}")

# Reopen and check
check = Document(OUTPUT)
full_text = ' '.join([p.text for p in check.paragraphs])
for t in check.tables:
    for r in t.rows:
        for c in r.cells:
            full_text += ' ' + c.text

required = [
    STUDENT,
    MSSV,
    'VAI TRÒ CỦA GATEWAY',
    'github.com/dinhvaren/iot-gateway-security-lab',
]
missing_checks = []
for item in required:
    if item in full_text:
        print(f"  [PASS] '{item}' found")
    else:
        print(f"  [FAIL] '{item}' NOT FOUND!")
        missing_checks.append(item)

placeholders_gone = '[NHẬP TÊN ĐỀ TÀI' not in full_text
print(f"  [{'PASS' if placeholders_gone else 'FAIL'}] Placeholder removed: {placeholders_gone}")

print(f"\n  Paragraphs: {len(check.paragraphs)}, Tables: {len(check.tables)}")
print(f"  DOCX REOPEN: PASS")

if missing_checks:
    print("\n*** WARNING: Some required items missing! ***")
    sys.exit(1)
else:
    print("\n*** ALL CHECKS PASSED ***")
